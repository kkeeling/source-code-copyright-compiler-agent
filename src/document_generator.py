import logging
import os
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

class DocumentGenerator:
    def __init__(self, project_name, project_root):
        self.project_name = project_name
        self.project_root = project_root
        self.document = Document()
        self._setup_document()
        logging.info(f"Initializing DocumentGenerator for project: {self.project_name}")

    def _setup_document(self):
        """Set up the document with basic styles and title."""
        self.document.add_heading(f"Source Code: {self.project_name}", 0)

        # Add styles
        styles = self.document.styles
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(6.5)

        # Add custom heading style for code snippets
        snippet_heading_style = styles.add_style('SnippetHeading', WD_STYLE_TYPE.PARAGRAPH)
        snippet_heading_style.font.size = Pt(10)
        snippet_heading_style.font.bold = True

    def add_code_snippet(self, file_path, code_snippet):
        """
        Add a code snippet to the document.

        Args:
            file_path (str): The path of the file.
            code_snippet (str): The extracted code snippet.
        """
        relative_path = os.path.relpath(file_path, self.project_root)
        heading = self.document.add_paragraph(relative_path, style='SnippetHeading')
        
        # Split the code snippet into lines
        lines = code_snippet.split('\n')
        
        # Get the first and last 10 "pages" (assuming 50 lines per page)
        first_pages = '\n'.join(lines[:500])
        last_pages = '\n'.join(lines[-500:])
        
        # Add the first 10 "pages"
        self.document.add_paragraph(first_pages, style='Code')
        
        # Add a separator if there's more content
        if len(lines) > 1000:
            self.document.add_paragraph("...\n[Content omitted for brevity]\n...")
        
        # Add the last 10 "pages" if they're different from the first 10
        if last_pages != first_pages:
            self.document.add_paragraph(last_pages, style='Code')

    def save_document(self, output_path):
        """
        Save the generated document.

        Args:
            output_path (str): The path to save the document.
        """
        self.document.save(output_path)
        logging.info(f"Document saved to: {output_path}")

if __name__ == "__main__":
    # Example usage
    doc_gen = DocumentGenerator("Test Project", "/path/to/project/root")
    doc_gen.add_code_snippet("/path/to/project/root/src/main.py", "print('Hello, World!')")
    doc_gen.save_document("test_output.docx")